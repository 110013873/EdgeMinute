"""无模型依赖的纯业务逻辑：时间格式化、发言人显示名、段落合并、
模板上下文构造、LLM 会议上下文构造。

刻意与 app.py 分离，使这些函数无需加载 FunASR 模型即可被单元测试覆盖。
所有函数均为纯函数，不修改入参（遵循不可变约定）。
"""
from __future__ import annotations

import io
import json
import math
import os
import re
from typing import Any


def _clean_env_value(raw: str) -> str:
    """清洗环境变量原始值，容忍常见的 .env 手写瑕疵。

    systemd 的 EnvironmentFile 不支持行尾 `#` 注释——`KEY=524288000  # 500MB`
    会把 `524288000  # 500MB` 整个当成值，导致 int()/float() 解析崩溃。
    这里剥离行尾注释、外围引号和首尾空白，使数值型配置更健壮。
    仅当 `#` 前存在空白时才视为行尾注释，避免误伤值内合法的 `#`。
    """
    s = str(raw).strip()
    m = re.search(r"\s#", s)
    if m:
        s = s[: m.start()].strip()
    if len(s) >= 2 and s[0] == s[-1] and s[0] in ("'", '"'):
        s = s[1:-1].strip()
    return s


def env_int(name: str, default: int) -> int:
    """读取整数型环境变量；缺失/空/无法解析时回退到 default（清洗行尾注释后再解析）。"""
    cleaned = _clean_env_value(os.environ.get(name, ""))
    if not cleaned:
        return default
    try:
        return int(cleaned)
    except (TypeError, ValueError):
        return default


def env_float(name: str, default: float) -> float:
    """读取浮点型环境变量；缺失/空/无法解析时回退到 default（清洗行尾注释后再解析）。"""
    cleaned = _clean_env_value(os.environ.get(name, ""))
    if not cleaned:
        return default
    try:
        return float(cleaned)
    except (TypeError, ValueError):
        return default


def format_time(seconds: float) -> str:
    """秒 → HH:MM:SS（向下取整）。"""
    s = int(seconds)
    h, m, sec = s // 3600, (s % 3600) // 60, s % 60
    return f"{h:02d}:{m:02d}:{sec:02d}"


def speaker_display(raw: Any, speaker_map: dict[str, str] | None = None) -> str:
    """发言人显示名：优先用 speaker_map；纯数字加“说话人”前缀；其余原样返回。

    与前端 speakerLabel 语义保持一致。
    """
    key = str(raw)
    if speaker_map:
        mapped = speaker_map.get(key, "")
        if mapped and mapped.strip():
            return mapped.strip()
    return f"说话人{key}" if key.isdigit() else key


def _join_text(a: str, b: str) -> str:
    """拼接两段文本：仅当衔接处两侧均为 ASCII 字母/数字时补空格。

    与前端 joinText 语义一致：中文之间、中文标点后不加空格。
    """
    a, b = (a or "").strip(), (b or "").strip()
    if not a:
        return b
    if not b:
        return a
    need_space = bool(re.search(r"[A-Za-z0-9]$", a)) and bool(re.match(r"[A-Za-z0-9]", b))
    return a + (" " if need_space else "") + b


# 参会人员分隔符：与前端 attendeeList 的 split 正则保持一致
_ATTENDEE_SPLIT = re.compile(r"[、,，;；\n\r\t ]+")
_ATTENDEE_KEYS = ("name", "unit", "title")


def _normalize_attendees(meta: dict) -> list[dict]:
    """把 meta.attendees 归一化为 [{name, unit, title}] 列表（纯函数，不改入参）。

    兼容三种输入：
    - list[dict]：保留 name/unit/title 并去空白
    - str（旧版本）：按分隔符拆成多个姓名，unit/title 留空
    - None/缺失：返回 []
    姓名为空的条目会被丢弃。
    """
    raw = (meta or {}).get("attendees")
    items: list[dict] = []
    if isinstance(raw, str):
        for name in _ATTENDEE_SPLIT.split(raw):
            name = name.strip()
            if name:
                items.append({"name": name, "unit": "", "title": ""})
    elif isinstance(raw, list):
        for a in raw:
            if not isinstance(a, dict):
                continue
            entry = {k: str(a.get(k, "") or "").strip() for k in _ATTENDEE_KEYS}
            if entry["name"]:
                items.append(entry)
    return items


def attendees_to_lines(items: list[dict]) -> str:
    """把结构化参会人渲染为多行字符串，每人一行：姓名（单位·职务）。

    单位/职务缺失时优雅省略：`姓名`、`姓名（单位）`、`姓名（·职务）`→`姓名（职务）`。
    """
    lines = []
    for a in items or []:
        name = (a.get("name") or "").strip()
        if not name:
            continue
        unit = (a.get("unit") or "").strip()
        title = (a.get("title") or "").strip()
        extra = "·".join(x for x in (unit, title) if x)
        lines.append(f"{name}（{extra}）" if extra else name)
    return "\n".join(lines)


# 无单位参会人归入的分组名
_NO_UNIT_GROUP = "其他"


def attendees_by_unit(items: list[dict]) -> str:
    """把结构化参会人按单位分组渲染为多行字符串：每单位一行 `单位：张三、李四`。

    - 单位按首次出现顺序排列，单位内成员按出现顺序保留、去重。
    - 无单位者归入“其他”分组，且仅当确有此类成员时才出现该行。
    - 无任何有效参会人时返回空串。
    """
    groups: list[str] = []          # 保序的单位名列表
    members: dict[str, list[str]] = {}
    for a in items or []:
        name = (a.get("name") or "").strip()
        if not name:
            continue
        unit = (a.get("unit") or "").strip() or _NO_UNIT_GROUP
        if unit not in members:
            members[unit] = []
            groups.append(unit)
        if name not in members[unit]:      # 同组内姓名去重
            members[unit].append(name)
    lines = [f"{unit}：{'、'.join(members[unit])}" for unit in groups]
    return "\n".join(lines)


def _speaker_unit_map(items: list[dict]) -> dict[str, str]:
    """由参会人列表构造 姓名→单位 映射，用于给发言人标注所属单位。

    姓名去空白后作键；重名以首次出现的单位为准；单位为空的条目跳过。
    """
    out: dict[str, str] = {}
    for a in items or []:
        name = (a.get("name") or "").strip()
        unit = (a.get("unit") or "").strip()
        if name and unit and name not in out:
            out[name] = unit
    return out


def build_speaker_blocks(
    segments: list[dict],
    speaker_map: dict[str, str] | None = None,
    attendees: list[dict] | None = None,
) -> list[dict]:
    """把段落按连续发言人合并成发言块，供“各发言人意见”分段展示。

    返回 [{speaker, unit, time, text}]：
    - speaker：经 speaker_map 解析后的显示名（如“张三”“说话人0”）。
    - unit：该发言人在参会人中的单位；匹配不到则为空串。
    - time：该发言块的时间范围 HH:MM:SS-HH:MM:SS（模板可按需展示）。
    - text：该发言人连续发言合并后的文本。
    连续同发言人段落合并为一块（复用 merge_segments 的相邻合并语义）。
    """
    speaker_map = speaker_map or {}
    unit_of = _speaker_unit_map(attendees or [])
    merged = merge_segments(segments or [], merge=True)
    blocks: list[dict] = []
    for s in merged:
        name = speaker_display(s.get("speaker", "?"), speaker_map)
        blocks.append({
            "speaker": name,
            "unit": unit_of.get(name, ""),
            "time": f"{format_time(s.get('start', 0))}-{format_time(s.get('end', 0))}",
            "text": (s.get("text") or "").strip(),
        })
    return blocks


def build_speaker_opinions(
    speakers: list[dict] | None,
    attendees: list[dict] | None = None,
) -> list[dict]:
    """由会议总结的 speakers（[{speaker, points}]）构造模板用的发言人观点列表。

    返回 [{speaker, unit, points}]：
    - speaker：发言人显示名（沿用总结中出现的名称，后端总结已按 speakerMap 映射）。
    - unit：该发言人在参会人中的单位；姓名匹配不到则为空串。
    - points：该发言人的核心观点总结（原样保留，去首尾空白）。
    仅保留 speaker/points 至少其一非空的条目。
    """
    unit_of = _speaker_unit_map(attendees or [])
    out: list[dict] = []
    for s in speakers or []:
        if not isinstance(s, dict):
            continue
        speaker = str(s.get("speaker", "") or "").strip()
        points = str(s.get("points", "") or "").strip()
        if not (speaker or points):
            continue
        out.append({
            "speaker": speaker,
            "unit": unit_of.get(speaker, ""),
            "points": points,
        })
    return out


# 决议条目行首序号，用于剥离 LLM 误加的前缀：1. / 1、/ 1) / （1）/ (1) / ①-⑳ / 一、二、 / - •
_RESOLUTION_NUM_PREFIX = re.compile(
    r"^\s*(?:"
    r"[（(]?\d+[)）]"          # (1) （1） 1)
    r"|\d+\s*[.、．]"           # 1. 1、 1．
    r"|[①-⑳]"                  # 带圈数字
    r"|[一二三四五六七八九十]+\s*[、.．]"  # 一、二、
    r"|[-•*]\s"                # - • * 列表符
    r")\s*"
)

# 行内序号断点：把挤在同一行的多条决议（如“1. 甲 2. 乙”）拆开
_RESOLUTION_INLINE_SPLIT = re.compile(
    r"(?<!^)\s+(?=(?:[（(]?\d+[)）]|\d+\s*[.、．]|[①-⑳]|[一二三四五六七八九十]+\s*[、.．]))"
)


def _strip_item_prefix(item: str) -> str:
    """去掉单条决议开头的序号/列表符（若有），返回条目正文。"""
    return _RESOLUTION_NUM_PREFIX.sub("", item.strip(), count=1).strip()


def normalize_resolutions(raw: Any) -> list[str]:
    """把会议决议规范为“每条一项”的字符串数组（纯函数）。

    兼容三种输入：
    - list：逐元素转字符串，去空白丢空项；元素内含多条（换行/行内序号）时再拆分。
    - str（旧数据 / LLM 未按数组返回）：按换行 + 行内序号拆条。
    - None/其它：返回 []。
    每条会剥离 LLM 误加的行首序号/列表符（展示与导出统一由前端/模板加序号或换行）。
    """
    raw_items: list[str]
    if isinstance(raw, list):
        raw_items = [str(x or "") for x in raw]
    elif isinstance(raw, str):
        raw_items = [raw]
    else:
        return []

    out: list[str] = []
    for chunk in raw_items:
        for line in chunk.splitlines():
            line = line.strip()
            if not line:
                continue
            for part in _RESOLUTION_INLINE_SPLIT.sub("\n", line).split("\n"):
                item = _strip_item_prefix(part)
                if item:
                    out.append(item)
    return out


def format_resolutions(raw: Any) -> str:
    """把会议决议规范为“按序号逐条换行”的字符串（供 docx 模板 {{ resolutions }}）。

    在 normalize_resolutions 的基础上重新编号为 1. 2. 3.，用 \\n 连接
    （docxtpl 会把 \\n 自动转 <w:br/>）。无决议时返回空串。
    """
    items = normalize_resolutions(raw)
    return "\n".join(f"{i}. {item}" for i, item in enumerate(items, 1))


def merge_segments(segments: list[dict], merge: bool = True) -> list[dict]:
    """合并连续的同发言人段落，返回全新列表（不修改入参）。

    merge=False 时仅返回逐段浅拷贝（text 去空白），不做合并。
    """
    if not merge:
        return [{**seg, "text": (seg.get("text") or "").strip()} for seg in segments]
    out: list[dict] = []
    for seg in segments:
        last = out[-1] if out else None
        if last is not None and str(last.get("speaker")) == str(seg.get("speaker")):
            out[-1] = {
                **last,
                "end": seg.get("end", last.get("end")),
                "text": _join_text(last.get("text", ""), seg.get("text", "")),
            }
        else:
            out.append({**seg, "text": (seg.get("text") or "").strip()})
    return out


# datetime-local 存储为 "YYYY-MM-DDTHH:mm"，展示时把 T 换成空格
def _display_date(raw: str) -> str:
    return (raw or "").replace("T", " ")


# 标量 meta 键（host、recorder 已移除；attendees 改为结构化，单独处理）
_META_KEYS = ("title", "date", "place", "agenda")


def build_template_context(
    meta: dict, segments: list[dict], options: dict, summary: dict | None = None
) -> dict:
    """构造 docxtpl 渲染上下文。

    options: {mergeSpeaker, timestamp, speaker, speakerMap}
    summary: {overview, speakers:[{speaker,points}], resolutions}
             （会议总结，可选；缺失时对应占位符渲染为空）
    - meta 缺失字段填空串；title 缺失回退为“会议纪要”
    - segments 依据 mergeSpeaker 合并，发言人经 speakerMap 解析显示名
    - 参会人员按单位分组（{{ attendees }}）
    - 各发言人意见取自 summary.speakers 的“核心观点总结”（{{ speaker_opinions }}，
      每项 {speaker, unit, points}），不再放原始转写内容
    - 会议概要 {{ overview }}；会议决议 {{ resolutions }}（已按序号逐条换行）
    - 兼容保留 {{ speaker_blocks }}（原始转写分段），供旧模板/回退使用
    """
    meta = meta or {}
    options = options or {}
    summary = summary or {}
    speaker_map = options.get("speakerMap") or {}
    attendees = _normalize_attendees(meta)

    merged = merge_segments(segments or [], merge=bool(options.get("mergeSpeaker", True)))
    seg_ctx = [
        {
            "time": f"{format_time(s.get('start', 0))}-{format_time(s.get('end', 0))}",
            "speaker": speaker_display(s.get("speaker", "?"), speaker_map),
            "text": s.get("text", ""),
        }
        for s in merged
    ]

    ctx: dict[str, Any] = {k: (meta.get(k) or "").strip() for k in _META_KEYS}
    ctx["date"] = _display_date(ctx["date"])
    if not ctx["title"]:
        ctx["title"] = "会议纪要"
    # 参会人员：结构化 → 按单位分组的多行字符串（{{ attendees }}）
    ctx["attendees"] = attendees_by_unit(attendees)
    ctx["show_speaker"] = bool(options.get("speaker", True))
    ctx["show_timestamp"] = bool(options.get("timestamp", False))
    ctx["segments"] = seg_ctx
    # 各发言人意见：取自会议总结的核心观点（{%p for b in speaker_opinions %}）
    ctx["speaker_opinions"] = build_speaker_opinions(summary.get("speakers"), attendees)
    # 兼容保留：原始转写分段（旧模板 {%p for b in speaker_blocks %} 仍可用）
    ctx["speaker_blocks"] = build_speaker_blocks(segments or [], speaker_map, attendees)
    # 会议概要 / 会议决议：来自会议总结（缺失时留空，模板占位符渲染为空串）
    ctx["overview"] = str(summary.get("overview", "") or "").strip()
    # 会议决议：数组 → 重新编号逐条换行（docxtpl 会把 \n 自动转 <w:br/>）；兼容旧字符串
    ctx["resolutions"] = format_resolutions(summary.get("resolutions"))
    return ctx


def build_meeting_context(
    meta: dict,
    segments: list[dict],
    budget: int = 24000,
    speaker_map: dict[str, str] | None = None,
) -> str:
    """把会议信息 + 转写文本拼装为注入 LLM 的上下文文本，按字符预算裁剪。

    speaker_map：发言人映射（如 {"0": "张三"}），使注入 LLM 的转写文本用
    映射后的姓名而非“说话人0”，否则用户按姓名提问时模型找不到对应发言。
    超预算时保留**开头**（会议通常前部信息密度高），并在末尾标注“已截断”。
    """
    meta = meta or {}
    speaker_map = speaker_map or {}
    header_parts = []
    # 标量字段（不含 attendees，参会人单独结构化渲染）
    label_map = {
        "title": "会议名称", "date": "时间", "place": "地点",
        "agenda": "议程",
    }
    for k, label in label_map.items():
        v = (meta.get(k) or "").strip()
        if v:
            header_parts.append(f"{label}：{_display_date(v) if k == 'date' else v}")
    attendees_text = attendees_to_lines(_normalize_attendees(meta))
    if attendees_text:
        header_parts.append("参会人员：\n" + attendees_text)
    header = "【会议信息】\n" + ("\n".join(header_parts) if header_parts else "（无）")

    merged = merge_segments(segments or [], merge=True)
    lines = [f"{speaker_display(s.get('speaker', '?'), speaker_map)}：{s.get('text', '')}" for s in merged]
    body = "\n".join(lines) if lines else "（暂无转写内容）"

    truncated = False
    if len(body) > budget:
        cut = body[:budget]
        # 尽量在行边界截断，避免切断一句话；找不到换行则硬截
        nl = cut.rfind("\n")
        body = cut[:nl] if nl > 0 else cut
        truncated = True

    parts = [header, "\n【会议记录】\n" + body]
    if truncated:
        parts.append("\n（注：会议记录较长，以上仅为按时间顺序保留的前段内容，后续部分已省略。）")
    return "\n".join(parts)


SYSTEM_PROMPT_TEMPLATE = (
    "你是一个会议助手。请依据下面提供的会议记录回答用户的问题，"
    "回答要准确、简洁，忠于会议内容；若会议记录中没有相关信息，请如实说明。\n\n{context}"
)


def build_llm_messages(payload: dict, budget: int = 24000) -> list[dict]:
    """把会议上下文注入 system，与前端传来的对话历史拼接。

    payload: { meta, segments, speakerMap, messages:[{role,content}] }
    仅保留 user/assistant 历史，避免前端注入额外 system 消息。
    speakerMap 用于把注入 LLM 的转写文本发言人替换为映射后的姓名。
    """
    meta = (payload or {}).get("meta") or {}
    segments = (payload or {}).get("segments") or []
    speaker_map = (payload or {}).get("speakerMap") or {}
    history = (payload or {}).get("messages") or []
    context = build_meeting_context(meta, segments, budget=budget, speaker_map=speaker_map)
    system = {"role": "system", "content": SYSTEM_PROMPT_TEMPLATE.format(context=context)}
    convo = [
        {"role": m["role"], "content": m["content"]}
        for m in history
        if m.get("role") in ("user", "assistant") and m.get("content")
    ]
    return [system, *convo]


# ---------- 会议总结：LLM 提示词 / 消息构造 / JSON 解析 ----------

_SUMMARY_SYSTEM_PROMPT = (
    "你是一个专业的会议纪要助手。用户会提供一场会议的完整记录（含发言人与发言内容）。"
    "请通读全部记录，撰写一份**详实、完整**的会议总结，并**只输出一个 JSON 对象**，"
    "不要输出任何解释、前后缀或 Markdown 代码块围栏。\n"
    "JSON 结构如下：\n"
    "{\n"
    '  "overview": "会议概要（字符串，300-500字，说明会议主题与背景、讨论的主要议题、各方达成的共识与分歧、核心结论）",\n'
    '  "speakers": [ { "speaker": "发言人姓名或标识", "points": "该发言人的观点详述（字符串）" } ],\n'
    '  "resolutions": [ "决议条目一", "决议条目二" ]\n'
    "}\n"
    "规则：\n"
    "- **speakers 是重点，务必详尽**：逐一覆盖每位有实质发言的发言人，不要遗漏。"
    "每位的 points 需完整归纳其观点，包括：提出的主张或立场、支撑的理由或依据、"
    "给出的建议或方案、提到的数据/事实/案例、表达的疑虑或反对意见、以及对他人观点的回应。"
    "应展开成多个要点（用换行分隔的分点文本，通常 3-6 点，内容多时可更多），"
    "而非一两句带过；宁可详实，避免过度精炼导致信息缺失。\n"
    "- 发言人标识沿用会议记录中出现的名称（如“张三”“说话人0”）。\n"
    "- resolutions 为字符串数组，每个元素是一条决议/待办（尽量含责任人、时间节点、具体事项），"
    "不要在元素内自带序号；无明确决议则返回只含一句说明的数组（如 [\"本次会议未形成明确决议\"]）。\n"
    "- 忠于会议记录，全面提炼记录中的实质信息，但不要编造记录中不存在的内容；"
    "记录为空或信息不足时如实说明。\n"
    "- overview/points 使用简洁清晰的中文；points 用换行分隔的分点文本。"
    "整体必须是可被 JSON.parse 直接解析的纯 JSON（换行在 JSON 字符串中写作 \\n）。"
)


def build_summary_messages(
    meta: dict,
    segments: list[dict],
    budget: int = 24000,
    speaker_map: dict[str, str] | None = None,
) -> list[dict]:
    """构造会议总结的 LLM 消息（system 约束 JSON 输出 + user 提供会议上下文）。

    复用 build_meeting_context 注入会议信息与转写文本，speaker_map 使发言人
    以映射后的姓名出现（如 0→张三），与问答 /export 行为一致。
    """
    context = build_meeting_context(
        meta or {}, segments or [], budget=budget, speaker_map=speaker_map or {}
    )
    user = context + "\n\n请依据以上会议内容进行总结，并按要求只输出 JSON。"
    return [
        {"role": "system", "content": _SUMMARY_SYSTEM_PROMPT},
        {"role": "user", "content": user},
    ]


def parse_summary_json(raw: str) -> dict:
    """解析 LLM 返回的总结 JSON，返回归一化后的结构。

    容错：去代码围栏、抽取首个平衡 {}、严格校验。解析失败抛 ValueError。
    返回：{overview:str, speakers:[{speaker,points}], resolutions:[str]}
    """
    try:
        obj = json.loads(_extract_json_object(raw), strict=False)
    except (ValueError, json.JSONDecodeError) as e:
        raise ValueError(f"LLM 返回无法解析为 JSON：{e}") from e
    if not isinstance(obj, dict):
        raise ValueError("LLM 返回的不是 JSON 对象")

    speakers: list[dict] = []
    raw_speakers = obj.get("speakers")
    if isinstance(raw_speakers, list):
        for s in raw_speakers:
            if not isinstance(s, dict):
                continue
            speaker = str(s.get("speaker", "") or "").strip()
            points = str(s.get("points", "") or "").strip()
            if speaker or points:
                speakers.append({"speaker": speaker, "points": points})

    return {
        "overview": str(obj.get("overview", "") or "").strip(),
        "speakers": speakers,
        "resolutions": normalize_resolutions(obj.get("resolutions")),
    }


# ---------- 议程导入：Word 文本提取 / LLM 提示词 / JSON 解析 ----------

AGENDA_TEXT_MAX_CHARS = 20000  # 送入 LLM 的文档文本上限


def extract_docx_text(data: bytes) -> str:
    """从 .docx 字节流提取纯文本（段落 + 表格单元格），供 LLM 提取字段。

    议程通知常用表格排版，故必须读取表格单元格文本，否则会漏掉参会人员等信息。
    解析失败抛 ValueError。
    """
    try:
        from docx import Document  # docxtpl 的传递依赖，无需额外安装
    except ImportError as e:  # pragma: no cover - 环境缺失时的兜底
        raise ValueError(f"缺少 python-docx 依赖：{e}") from e
    try:
        doc = Document(io.BytesIO(data))
    except Exception as e:
        raise ValueError(f"无法读取 Word 文档（请确认为 .docx 格式）：{e}") from e

    lines: list[str] = []
    for p in doc.paragraphs:
        t = (p.text or "").strip()
        if t:
            lines.append(t)
    for table in doc.tables:
        for row in table.rows:
            cells = [(c.text or "").strip() for c in row.cells]
            cells = [c for c in cells if c]
            if cells:
                lines.append(" | ".join(cells))

    text = "\n".join(lines).strip()
    if not text:
        raise ValueError("文档内容为空，未能提取到文本")
    return text[:AGENDA_TEXT_MAX_CHARS]


_AGENDA_SYSTEM_PROMPT = (
    "你是一个会议信息提取助手。用户会给你一份“会议议程通知”文档的纯文本，"
    "请从中提取会议关键信息，并**只输出一个 JSON 对象**，不要输出任何解释、前后缀或 Markdown 代码块围栏。\n"
    "JSON 结构如下：\n"
    "{\n"
    '  "title": "会议名称（字符串）",\n'
    '  "date": "会议时间（字符串，尽量保留原文，如 2026-07-30 14:00）",\n'
    '  "place": "会议地点（字符串）",\n'
    '  "agenda": "会议议程（字符串，多项议程用换行分隔）",\n'
    '  "attendees": [ { "name": "姓名", "unit": "单位", "title": "职务" } ]\n'
    "}\n"
    "规则：\n"
    "- 未提及的字段填空字符串 \"\"，attendees 未提及时填空数组 []。\n"
    "- attendees 中每个人的 unit/title 若文档未写明则填空字符串。\n"
    "- 不要编造文档中不存在的信息。\n"
    "- 再次强调：输出必须是可被 JSON.parse 直接解析的纯 JSON，不含任何额外文字。"
)


def build_agenda_extract_messages(text: str) -> list[dict]:
    """构造议程提取的 LLM 消息（system 约束 JSON 输出 + user 提供文档文本）。"""
    user = "以下是会议议程通知文档的文本内容：\n---\n" + (text or "") + "\n---\n请提取并输出 JSON。"
    return [
        {"role": "system", "content": _AGENDA_SYSTEM_PROMPT},
        {"role": "user", "content": user},
    ]


def _extract_json_object(raw: str) -> str:
    """从可能包含围栏/前后缀说明的 LLM 输出中抽出 JSON 对象子串。"""
    s = (raw or "").strip()
    # 去 ```json ... ``` 或 ``` ... ``` 围栏
    fence = re.search(r"```(?:json)?\s*(.+?)\s*```", s, re.DOTALL)
    if fence:
        s = fence.group(1).strip()
    if s.startswith("{") and s.endswith("}"):
        return s
    # 扫描第一个平衡的 {...}，字符串字面量内的花括号不计入深度
    start = s.find("{")
    if start == -1:
        raise ValueError("LLM 返回中未找到 JSON 对象")
    depth = 0
    in_string = False
    escaped = False
    for i in range(start, len(s)):
        ch = s[i]
        if in_string:
            if escaped:
                escaped = False
            elif ch == "\\":
                escaped = True
            elif ch == '"':
                in_string = False
            continue
        if ch == '"':
            in_string = True
        elif ch == "{":
            depth += 1
        elif ch == "}":
            depth -= 1
            if depth == 0:
                return s[start:i + 1]
    raise ValueError("LLM 返回的 JSON 对象不完整")


def parse_agenda_json(raw: str) -> dict:
    """解析 LLM 返回的议程 JSON，返回归一化后的 meta 结构。

    容错：去代码围栏、抽取首个平衡 {}、严格校验。解析失败抛 ValueError。
    返回：{title, date, place, agenda, attendees:[{name,unit,title}]}
    """
    try:
        # strict=False 容忍 LLM 在 JSON 字符串里直接输出的原始换行等控制字符
        obj = json.loads(_extract_json_object(raw), strict=False)
    except (ValueError, json.JSONDecodeError) as e:
        raise ValueError(f"LLM 返回无法解析为 JSON：{e}") from e
    if not isinstance(obj, dict):
        raise ValueError("LLM 返回的不是 JSON 对象")

    meta = {
        "title": str(obj.get("title", "") or "").strip(),
        "date": str(obj.get("date", "") or "").strip(),
        "place": str(obj.get("place", "") or "").strip(),
        "agenda": str(obj.get("agenda", "") or "").strip(),
    }
    # 复用统一的参会人归一化（同时兼容对象数组或误返回的字符串）
    meta["attendees"] = _normalize_attendees(obj)
    return meta


# ---------- 会议历史：payload 构造 / 冗余列抽取 / 校验 / 按日期分组 ----------

HISTORY_SCHEMA_VERSION = 2  # payload 结构版本，供将来迁移兼容
HISTORY_SNIPPET_CHARS = 120  # 列表页摘要字符数
_HISTORY_FILE_KEYS = ("id", "name", "duration", "elapsed", "audio_name", "has_audio")

# 会议级状态（DB 冗余列 + payload.status）
MEETING_DRAFT = "draft"            # 已建档，尚无任何转写结果
MEETING_TRANSCRIBING = "transcribing"  # 至少一个文件排队/转写中
MEETING_DONE = "done"             # 无进行中转写（全部完成或本就无需转写）

# 文件级状态（仅存于 payload.files[]）
_FILE_ACTIVE = {"queued", "processing"}  # 视为“转写进行中”的文件态


def derive_meeting_status(files: list[dict]) -> str:
    """由文件级 status 汇总出会议级 status（纯函数）。

    - 任一文件处于 queued/processing → transcribing
    - 否则若存在任何已完成的转写结果（有 segments 的文件）→ done
    - 否则（无文件，或文件均无结果且未在转写）→ draft
    兼容旧 payload：文件无 status 字段时按“有无 segments”判定完成/草稿。
    """
    files = files or []
    if any((f.get("status") in _FILE_ACTIVE) for f in files):
        return MEETING_TRANSCRIBING
    has_result = any((f.get("segments") or []) for f in files)
    return MEETING_DONE if has_result else MEETING_DRAFT


def _normalize_history_file(raw: Any) -> dict:
    """把单个文件条目归一化为 payload 内的 file 结构（纯函数，不改入参）。

    保留 id/name/duration/elapsed/audio_name/has_audio 与 segments（逐段浅拷贝）。
    音频引用 audio_name 为服务端相对文件名；has_audio 缺失时按 audio_name 是否存在推断。
    """
    raw = raw or {}
    segs_in = raw.get("segments") or []
    segments = []
    for s in segs_in:
        if not isinstance(s, dict):
            continue
        segments.append({
            "start": float(s.get("start", 0) or 0),
            "end": float(s.get("end", 0) or 0),
            "speaker": s.get("speaker", "?"),
            "text": str(s.get("text", "") or ""),
            "flagged": bool(s.get("flagged", False)),
        })
    audio_name = str(raw.get("audio_name", "") or "").strip()
    has_audio = raw.get("has_audio")
    # 文件级状态：pending=待转写 / queued=排队 / processing=转写中 / done / error。
    # 兼容旧 payload（无 status）：有 segments 视为 done，否则 pending。
    status = str(raw.get("status", "") or "").strip()
    if not status:
        status = "done" if segments else "pending"
    return {
        "id": str(raw.get("id", "") or ""),
        "name": str(raw.get("name", "") or ""),
        "duration": raw.get("duration"),
        "elapsed": raw.get("elapsed"),
        "audio_name": audio_name,
        "has_audio": bool(has_audio) if has_audio is not None else bool(audio_name),
        "status": status,
        "segments": segments,
    }


def build_history_payload(
    meta: dict,
    files: list[dict],
    speaker_map: dict[str, str] | None = None,
    summary: dict | None = None,
    status: str | None = None,
) -> dict:
    """组装一条会议历史的完整 payload（纯函数）。

    - meta 经参会人归一化，剔除非法条目。
    - files 每项保留元信息 + 音频引用 + segments + 文件级 status。
    - status：会议级状态；缺省由 derive_meeting_status(files) 推导。
      显式传入用于建"草稿档"（无文件也标 draft）等场景。
    - schema_version 内嵌，供将来迁移兼容。
    """
    meta = meta or {}
    norm_meta = {
        "title": str(meta.get("title", "") or "").strip(),
        "date": str(meta.get("date", "") or "").strip(),
        "place": str(meta.get("place", "") or "").strip(),
        "agenda": str(meta.get("agenda", "") or "").strip(),
        "attendees": _normalize_attendees(meta),
    }
    norm_files = [_normalize_history_file(f) for f in (files or [])]
    meeting_status = (status or "").strip() or derive_meeting_status(norm_files)
    return {
        "schema_version": HISTORY_SCHEMA_VERSION,
        "meta": norm_meta,
        "speakerMap": dict(speaker_map or {}),
        "summary": summary or {},
        "status": meeting_status,
        "files": norm_files,
    }


def extract_history_columns(payload: dict) -> dict:
    """从 payload 计算入库的冗余列（纯函数）。

    返回 {title, meeting_date, duration, segment_count, speaker_count, snippet}：
    - duration：各文件 duration 求和（None 记 0）。
    - segment_count：全部文件 segments 总数。
    - speaker_count：去重发言人数（按经映射解析后的显示名去重，与展示一致）。
    - snippet：合并后转写文本前 N 字（供列表摘要）。
    """
    payload = payload or {}
    meta = payload.get("meta") or {}
    files = payload.get("files") or []
    speaker_map = payload.get("speakerMap") or {}

    duration = 0.0
    all_segments: list[dict] = []
    for f in files:
        d = f.get("duration")
        if isinstance(d, (int, float)):
            duration += float(d)
        all_segments.extend(f.get("segments") or [])

    speakers = {
        speaker_display(s.get("speaker", "?"), speaker_map)
        for s in all_segments
    }
    merged = merge_segments(all_segments, merge=True)
    snippet = " ".join((s.get("text") or "").strip() for s in merged).strip()
    snippet = snippet[:HISTORY_SNIPPET_CHARS]

    # 会议级 status 冗余列：优先用 payload 显式值，否则由文件态推导（兼容旧数据）
    status = str(payload.get("status", "") or "").strip() or derive_meeting_status(files)

    return {
        "title": str(meta.get("title", "") or "").strip(),
        "meeting_date": str(meta.get("date", "") or "").strip(),
        "duration": round(duration, 1),
        "segment_count": len(all_segments),
        "speaker_count": len(speakers),
        "snippet": snippet,
        "status": status,
    }


def validate_history_payload(raw: Any, max_bytes: int | None = None) -> dict:
    """校验并归一化外来 payload（入库前的边界防护）。

    - 接受 dict 或 JSON 字符串；结构非法抛 ValueError。
    - 超过 max_bytes（若给出，按 UTF-8 JSON 字节计）抛 ValueError。
    - 返回经 build_history_payload 重建的规范 payload（丢弃未知字段）。
    """
    if isinstance(raw, str):
        try:
            raw = json.loads(raw, strict=False)
        except (ValueError, json.JSONDecodeError) as e:
            raise ValueError(f"payload 不是合法 JSON：{e}") from e
    if not isinstance(raw, dict):
        raise ValueError("payload 必须是 JSON 对象")

    files = raw.get("files")
    if files is not None and not isinstance(files, list):
        raise ValueError("payload.files 必须是数组")

    payload = build_history_payload(
        meta=raw.get("meta") or {},
        files=files or [],
        speaker_map=raw.get("speakerMap") or {},
        summary=raw.get("summary") or {},
        status=raw.get("status"),
    )
    if max_bytes is not None:
        size = len(json.dumps(payload, ensure_ascii=False).encode("utf-8"))
        if size > max_bytes:
            raise ValueError(f"payload 过大（{size} 字节，上限 {max_bytes}）")
    return payload


def group_by_date(items: list[dict]) -> list[dict]:
    """把列表项按 meeting_date 的“天”分组（纯函数，供前端亦可复用的分组逻辑参考）。

    - 分组键取 meeting_date 的日期部分（"YYYY-MM-DDTHH:mm" / "YYYY-MM-DD ..." → "YYYY-MM-DD"）。
    - 无法解析日期的归入 "未标注日期" 组，排在最后。
    - 组按日期倒序；组内保持传入顺序（调用方应已按时间倒序）。
    返回 [{date, items:[...]}]。
    """
    UNKNOWN = "未标注日期"
    groups: dict[str, list[dict]] = {}
    order: list[str] = []
    for it in items or []:
        raw = str((it or {}).get("meeting_date", "") or "").strip()
        m = re.match(r"(\d{4}-\d{2}-\d{2})", raw.replace("T", " "))
        key = m.group(1) if m else UNKNOWN
        if key not in groups:
            groups[key] = []
            order.append(key)
        groups[key].append(it)
    # 已知日期倒序，未知组始终置尾
    known = sorted((k for k in order if k != UNKNOWN), reverse=True)
    ordered = known + ([UNKNOWN] if UNKNOWN in groups else [])
    return [{"date": k, "items": groups[k]} for k in ordered]


# ---------- 声纹匹配：向量归一化 / 余弦相似度 / 说话人匹配（纯函数） ----------
# 刻意用纯 Python 数学实现，使这些函数无需 numpy 即可单测（与本模块“无重依赖”一致）。
# 向量入参接受任意可迭代的数值序列（list / numpy 一维数组皆可）。


def l2_normalize(vec: Any) -> list[float]:
    """L2 归一化一维向量，返回新的 list[float]（不修改入参）。

    零向量（或近零）原样返回全 0，避免除零；配合余弦：归一化后点积即余弦相似度。
    """
    v = [float(x) for x in (vec or [])]
    norm = math.sqrt(sum(x * x for x in v))
    if norm <= 1e-12:
        return [0.0 for _ in v]
    return [x / norm for x in v]


def cosine_similarity(a: Any, b: Any) -> float:
    """两向量余弦相似度 ∈ [-1, 1]。任一为零向量或维度不匹配时返回 0.0。

    不假设入参已归一化：内部各自除以模长，故可直接传原始声纹向量。
    """
    va = [float(x) for x in (a or [])]
    vb = [float(x) for x in (b or [])]
    if not va or not vb or len(va) != len(vb):
        return 0.0
    na = math.sqrt(sum(x * x for x in va))
    nb = math.sqrt(sum(x * x for x in vb))
    if na <= 1e-12 or nb <= 1e-12:
        return 0.0
    dot = sum(x * y for x, y in zip(va, vb))
    return dot / (na * nb)


def match_speaker(embedding: Any, enrolled: list[dict], threshold: float) -> dict:
    """把单个说话人簇的声纹向量匹配到登记库中最相似的人。

    - embedding：该说话人簇的代表性声纹向量。
    - enrolled：登记库 [{id, name, embedding:[...]}, ...]（embedding 为一维数值序列）。
    - threshold：余弦相似度阈值；最高分 ≥ 阈值才算命中。
    返回 {matched:bool, id, name, score}；未命中时 matched=False、id/name 为 None，
    score 仍给出最高分（便于日志/调阈值）。库为空或向量非法时 matched=False、score=0。
    """
    best_id = None
    best_name = None
    best_score = -1.0
    for person in enrolled or []:
        if not isinstance(person, dict):
            continue
        score = cosine_similarity(embedding, person.get("embedding"))
        if score > best_score:
            best_score = score
            best_id = person.get("id")
            best_name = person.get("name")
    if best_score < 0:  # 库为空（未进入循环）：无最高分
        return {"matched": False, "id": None, "name": None, "score": 0.0}
    matched = best_score >= threshold
    return {
        "matched": matched,
        "id": best_id if matched else None,
        "name": best_name if matched else None,
        "score": best_score,
    }


def match_speakers(
    cluster_embeddings: dict, enrolled: list[dict], threshold: float
) -> dict:
    """把多个说话人簇一次性匹配到登记库，供转写后填充 speakerMap。

    - cluster_embeddings：{原始 spk 值(str): 该簇声纹向量}。
    - enrolled / threshold：同 match_speaker。
    返回 {spk(str): {name, score}}，**仅包含命中的簇**（未命中的不出现，
    使前端可原样保留“说话人N”）。姓名冲突不去重——同一人匹配到多个簇属正常
    （聚类可能把一个人拆成两簇），由前端/用户决定是否合并。
    """
    out: dict[str, dict] = {}
    for spk, emb in (cluster_embeddings or {}).items():
        res = match_speaker(emb, enrolled, threshold)
        if res["matched"]:
            out[str(spk)] = {"name": res["name"], "score": round(res["score"], 3)}
    return out


def build_representative_segments(
    sentence_info: list,
    max_segments: int,
    min_seconds: float,
) -> dict:
    """把 FunASR 的 sentence_info 按说话人聚类编号分组，各取**最长的若干段**，
    作为提取该簇代表性声纹的时间区间。纯函数（无音频/模型依赖，便于单测）。

    - sentence_info：[{start(ms), end(ms), spk, ...}]，start/end 为**毫秒**（原始单位）。
    - max_segments：每个说话人最多取几段（按时长降序）。
    - min_seconds：短于此秒数的段直接丢弃（太短向量不可靠）。

    返回 {spk(int/原值): [(start_sec, end_sec), ...]}，时间已换算成**秒**并按起点升序。
    某说话人所有段都过短则不出现在结果里（该簇无法可靠提向量）。
    """
    by_spk: dict = {}
    for seg in sentence_info or []:
        spk = seg.get("spk")
        st, en = seg.get("start"), seg.get("end")
        if spk is None or st is None or en is None:
            continue
        start_sec = st / 1000.0
        end_sec = en / 1000.0
        if end_sec - start_sec < min_seconds:
            continue
        by_spk.setdefault(spk, []).append((start_sec, end_sec))

    out: dict = {}
    for spk, spans in by_spk.items():
        top = sorted(spans, key=lambda p: p[1] - p[0], reverse=True)[: max(1, max_segments)]
        out[spk] = sorted(top, key=lambda p: p[0])
    return out
